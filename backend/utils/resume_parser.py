"""
HireSense-AI Pro — Resume Text Extractor
Supports PDF and DOCX files.
"""
import io
from typing import Union


def extract_resume_text(file_bytes: bytes, filename: str) -> str:
    """
    Extract plain text from a PDF or DOCX resume file.

    Args:
        file_bytes: Raw file bytes.
        filename:   Original filename (used to detect file type).

    Returns:
        Extracted plain text string.

    Raises:
        ValueError: If the file type is unsupported or text extraction fails.
    """
    filename_lower = filename.lower()

    if filename_lower.endswith(".pdf"):
        return _extract_from_pdf(file_bytes)
    elif filename_lower.endswith(".docx"):
        return _extract_from_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file type. Please upload a PDF or DOCX file.")


def _extract_from_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF using pdfplumber (best for ATS-style PDFs)."""
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        text = "\n".join(text_parts).strip()
        if not text:
            raise ValueError("PDF appears to be image-based. Please use a text-based PDF.")
        return text
    except ImportError:
        # Fallback to pypdf
        try:
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(file_bytes))
            text = "\n".join(
                page.extract_text() or "" for page in reader.pages
            ).strip()
            if not text:
                raise ValueError("Could not extract text from PDF.")
            return text
        except ImportError:
            raise ValueError("PDF library not installed. Run: pip install pdfplumber")


def _extract_from_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs).strip()
    except ImportError:
        raise ValueError("python-docx not installed. Run: pip install python-docx")
